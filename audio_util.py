# -*- coding: utf-8 -*-
"""
MedSnap 音频处理与文本分析工具模块
提供语音转写、结构化提取、质性研究分析及文本文件解析功能。

ASR 策略：
  使用 Qwen2.5-Omni 全模态模型直接处理音频文件，完成语音识别（ASR）。
  音频以 Base64 编码通过 OpenAI 兼容接口发送，无需 DashScope 依赖。
"""

import os
import re

import model
from model import parse_ai_response, audio_to_base64
from desensitizer import desensitize_text

# ========== 可选依赖检测（格式转换用） ==========

HAS_PYDUB = False
try:
    from pydub import AudioSegment
    HAS_PYDUB = True
except ImportError:
    pass

# ========== 音频格式 MIME 映射 ==========

AUDIO_MIME_MAP = {
    '.wav':  'audio/wav',
    '.mp3':  'audio/mpeg',
    '.aac':  'audio/aac',
    '.amr':  'audio/amr',
    '.opus': 'audio/opus',
    '.m4a':  'audio/mp4',
    '.flac': 'audio/flac',
}

# ========== ASR Prompt ==========

PROMPT_ASR = """请对以下音频内容进行语音识别，输出完整的转录文本。
要求：
1. 忠实还原说话人的原始表达，不要改写或总结
2. 中英文混合时保持原样
3. 只输出转录文本，不要输出任何其他内容"""

# ========== 质性分析 Prompt 模板 ==========

PROMPT_QUALITATIVE_ANALYSIS = """你是临床定性研究专家。请对以下医疗访谈转录文本进行定性分析。

输出JSON格式:
{{
  "themes": ["主题1", "主题2"],
  "keywords": ["关键词1", "关键词2"],
  "codes": [{{"code": "编码类别", "segments": ["相关文本片段1", "片段2"]}}],
  "sentiment": "积极/中性/消极",
  "summary": "2-3句话分析总结"
}}

分析要求:
1. 主题分析: 识别3-5个核心讨论主题
2. 关键词提取: 提取10-15个关键医学/情感词汇
3. 编码分类: 按类别(如症状描述、治疗态度、医患沟通、情感表达、生活影响)编码文本
4. 情感倾向: 判断整体情感

转录文本:
{transcript}

只输出JSON，不要输出任何其他内容。"""

QUALITATIVE_TYPE_HINTS = {
    'interview': {
        'cn': '深度访谈',
        'hints': '患者体验、疾病认知、治疗态度、就医过程、心理感受',
        'coding': '症状描述、情感表达、医患沟通、治疗依从性、生活影响'
    },
    'focus_group': {
        'cn': '焦点小组',
        'hints': '群体共识、争议点、互动模式、观点演变、关键事件',
        'coding': '观点类别、互动类型、共识差异、关键事件、群体动态'
    },
    'observation': {
        'cn': '观察记录',
        'hints': '行为模式、环境因素、非语言信息、事件序列、场景特征',
        'coding': '行为类型、场景因素、时间特征、主体角色、环境条件'
    }
}

PROMPT_QUALITATIVE_ENHANCED = """你是质性研究方法论专家。请对以下{analysis_type_cn}材料严格按照四步质性分析法进行系统分析。

## 分析步骤

### 第一步：初始编码（Open Coding）
逐句或逐段阅读文本，识别有意义的概念、想法、行为模式，并为每个有意义的片段打上编码标签。
- 每个编码包含唯一编号(C01, C02...)、编码标签（简短概念名）、对应的原文片段、段落编号(P1, P2...)

### 第二步：主题聚类（Theme Clustering）
将相似的初始编码进行归纳合并，形成更高层次的上位主题类别。
- 识别3-5个主主题
- 每个主主题下有1-4个子主题
- 每个子主题关联具体的code_id列表
- 确保主题内部逻辑一致性和主题间差异性

### 第三步：典型原话保留（Representative Quotes）
为每个最终确定的主题选择2-3条最具代表性的原始引语。
- 必须是原文直接引用，不做改编
- 应能充分支撑该主题的核心观点

### 第四步：层级化输出（Hierarchical Structure）
最终输出严格按照"主题—子主题—编码—原话摘录"的层级结构。

## 输出JSON格式（严格遵守此结构）
{{
  "methodology_note": "本分析采用{analysis_type_cn}质性研究方法，遵循开放性编码→主题聚类→代表性引用→层级输出的四步分析流程",
  "analysis_type": "{analysis_type}",
  "step1_initial_coding": [
    {{
      "code_id": "C01",
      "code_label": "编码标签名",
      "original_text": "原文中的具体片段",
      "paragraph_ref": "P1"
    }}
  ],
  "step2_theme_clustering": [
    {{
      "theme": "主主题名称",
      "sub_themes": [
        {{
          "sub_theme": "子主题名称",
          "codes": ["C01", "C03"],
          "description": "该子主题的简要描述"
        }}
      ]
    }}
  ],
  "step3_representative_quotes": [
    {{
      "theme": "主主题名称",
      "quotes": [
        "原话引用1",
        "原话引用2"
      ]
    }}
  ],
  "step4_hierarchical_output": [
    {{
      "theme": "主主题名称",
      "sub_themes": [
        {{
          "name": "子主题名称",
          "codes": [
            {{
              "label": "编码标签",
              "quotes": ["支撑该编码的原话"]
            }}
          ]
        }}
      ]
    }}
  ]
}}

## 分析要点
- 分析关注点：{analysis_hints}
- 编码参考类别：{coding_categories}
- 初始编码数量：8-20个（视文本长度而定）
- 主主题数量：3-5个
- 每个主题下的代表性引用：2-3条
- 所有引用必须来自原文，保持原始措辞

## 待分析文本：
{transcript}

只输出JSON，不要输出任何其他内容。"""

# ========== 语音识别 ==========

def transcribe_audio(audio_path):
    """
    使用 Qwen2.5-Omni 全模态模型对本地音频文件进行语音识别（ASR）。
    音频以 Base64 编码通过 OpenAI 兼容接口发送，无需 DashScope 依赖。

    支持格式：wav / mp3 / aac / amr / opus / m4a / flac
    不支持的格式会尝试用 pydub 转为 wav 后再识别。

    Args:
        audio_path: 音频文件路径

    Returns:
        dict: {'text': 完整转录文本, 'sentences': [], 'language': 'zh'}

    Raises:
        Exception: 文件不存在或识别失败时抛出
    """
    ext = os.path.splitext(audio_path)[1].lower()
    actual_path = audio_path
    converted = False

    # 不在 MIME 映射中的格式，尝试用 pydub 转为 wav
    if ext not in AUDIO_MIME_MAP:
        if not HAS_PYDUB:
            raise Exception(
                f"不支持的音频格式 {ext}，请安装 pydub 进行格式转换: pip install pydub"
            )
        audio_segment = AudioSegment.from_file(audio_path)
        actual_path = audio_path + '.wav'
        audio_segment.export(actual_path, format='wav')
        converted = True
        ext = '.wav'

    mime_type = AUDIO_MIME_MAP.get(ext, 'audio/wav')

    try:
        print(f"[ASR] Qwen2.5-Omni 语音识别：{os.path.basename(actual_path)}")
        b64_audio = audio_to_base64(actual_path)

        response = model.client.chat.completions.create(
            model=model.MODEL_NAME_OMNI,
            messages=[{
                'role': 'user',
                'content': [
                    {'type': 'text', 'text': PROMPT_ASR},
                    {
                        'type': 'input_audio',
                        'input_audio': {
                            'data': b64_audio,
                            'format': mime_type,
                        }
                    }
                ]
            }],
            temperature=0.1,
            max_tokens=4096
        )

        full_text = response.choices[0].message.content.strip()
        if not full_text:
            raise Exception("语音识别未返回有效文本，请检查音频文件质量")

        return {'text': full_text, 'sentences': [], 'language': 'zh'}
    finally:
        if converted and os.path.exists(actual_path):
            try:
                os.remove(actual_path)
            except Exception:
                pass


def extract_from_transcript(transcript_text, ai_prompt):
    """
    用 Qwen 模型从转录文本中提取结构化数据（纯文本模式，发送前自动脱敏）。

    Args:
        transcript_text: 转录文本
        ai_prompt: 提取用的 Prompt

    Returns:
        tuple: (parsed_dict, raw_text)
    """
    masked_text, _report = desensitize_text(transcript_text)
    combined_prompt = (
        ai_prompt
        + "\n\n以下是语音转录文本，请按上述要求提取结构化信息：\n\n"
        + masked_text
    )
    response = model.client.chat.completions.create(
        model=model.MODEL_NAME,
        messages=[{'role': 'user', 'content': combined_prompt}],
        temperature=0.1,
        max_tokens=4096
    )
    raw_text = response.choices[0].message.content
    parsed = parse_ai_response(raw_text)
    return parsed, raw_text


# ========== 质性研究分析 ==========

def qualitative_analysis(transcript_text):
    """
    对转录文本进行简版定性研究分析（发送前自动脱敏）。

    Args:
        transcript_text: 转录文本

    Returns:
        dict: 包含 themes、keywords、codes、sentiment、summary 的分析结果
    """
    masked_text, _report = desensitize_text(transcript_text)
    prompt = PROMPT_QUALITATIVE_ANALYSIS.format(transcript=masked_text)
    response = model.client.chat.completions.create(
        model=model.MODEL_NAME,
        messages=[{'role': 'user', 'content': prompt}],
        temperature=0.3,
        max_tokens=3000
    )
    return parse_ai_response(response.choices[0].message.content)


def qualitative_analysis_enhanced(transcript_text, analysis_type='interview'):
    """
    四步法质性分析：初始编码 → 主题聚类 → 代表性引用 → 层级输出（发送前自动脱敏）。

    Args:
        transcript_text: 转录文本
        analysis_type: 分析类型（interview / focus_group / observation）

    Returns:
        dict: 四步质性分析结果
    """
    masked_text, _report = desensitize_text(transcript_text)
    type_info = QUALITATIVE_TYPE_HINTS.get(analysis_type, QUALITATIVE_TYPE_HINTS['interview'])
    prompt = PROMPT_QUALITATIVE_ENHANCED.format(
        analysis_type_cn=type_info['cn'],
        analysis_type=analysis_type,
        analysis_hints=type_info['hints'],
        coding_categories=type_info['coding'],
        transcript=masked_text
    )
    response = model.client.chat.completions.create(
        model=model.MODEL_NAME,
        messages=[{'role': 'user', 'content': prompt}],
        temperature=0.3,
        max_tokens=4000
    )
    result = parse_ai_response(response.choices[0].message.content)
    # 兼容旧格式返回
    if 'error' not in result and 'step1_initial_coding' not in result:
        if 'themes' in result:
            return result
    return result


# ========== 文本文件解析 ==========

def parse_text_file(file_path):
    """
    解析 txt / docx 文件为纯文本字符串。

    Args:
        file_path: 文件路径

    Returns:
        str: 文件纯文本内容

    Raises:
        Exception: 不支持的格式或编码无法识别时抛出
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.txt':
        for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as text_file:
                    return text_file.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise Exception("无法识别文本文件编码，请使用UTF-8编码保存")

    elif ext == '.docx':
        try:
            import docx
        except ImportError:
            raise Exception("python-docx 库未安装，请运行 pip install python-docx")
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        return '\n'.join(text_parts)

    elif ext == '.doc':
        raise Exception("不支持旧版.doc格式，请将文件另存为.docx后重新上传")

    else:
        raise Exception(f"不支持的文本格式: {ext}")


def preprocess_text(text):
    """
    文本预处理：标准化换行、合并多余空白、去除首尾空格。

    Args:
        text: 原始文本

    Returns:
        str: 标准化后的文本
    """
    if not text:
        return ''
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()
